
import pandas as pd
import win32com.client as win32
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,Image
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import HRFlowable
from datetime import datetime


# Lecture du fichier Excel
df = pd.read_excel("back.xlsx", sheet_name="Demandes")
df = df.dropna(subset=["email"])
df["email"] = df["email"].str.strip()
df["nom_famille"] = df["nom_famille"].str.strip()
df["nom_prenom_eleve"] = df["nom_prenom_eleve"].str.strip()
df["ecole_sollicitee"] = df["ecole_sollicitee"].str.strip()
df["classe_sollicitee"] = df["classe_sollicitee"].str.strip()
df["decision"] = df["decision"].str.strip()
df["ref_demande"] = df["ref_demande"].str.strip()
df["Matricule"] = df["Matricule"].str.strip()

# Fonction construction tableau HTML
def construire_tableau_html(groupe):
    lignes = ""
    for i, row in enumerate(groupe.itertuples()):
        bg = "#ffffff" if i % 2 == 0 else "#f7f9fc"

        if str(row.decision).strip().lower() == "admis":
            couleur_decision = "#1a7a2e"
            texte_decision   = "Admis"
        else:
            couleur_decision = "#b4090c"
            texte_decision   = "Réfusé"

        lignes += f"""
        <tr style='background-color:{bg};'>
            <td style='padding:8px 10px;border:1px solid #c8d0dc;font-family:Arial,sans-serif;font-size:13px;color:#333;text-align:center;'>{i + 1}</td>
            <td style='padding:8px 10px;border:1px solid #c8d0dc;font-family:Arial,sans-serif;font-size:13px;color:#1a1a2e;font-weight:bold;'>{row.nom_prenom_eleve}</td>
            <td style='padding:8px 10px;border:1px solid #c8d0dc;font-family:Arial,sans-serif;font-size:13px;color:#333;text-align:center;'>{row.classe_sollicitee}</td>
            <td style='padding:8px 10px;border:1px solid #c8d0dc;font-family:Arial,sans-serif;font-size:13px;color:#333;text-align:center;'>{row.ecole_sollicitee}</td>
            <td style='padding:8px 10px;border:1px solid #c8d0dc;font-family:Arial,sans-serif;font-size:13px;color:#333;'>{row.ref_demande}</td>
            <td style='padding:8px 10px;border:1px solid #c8d0dc;font-family:Arial,sans-serif;font-size:13px;font-weight:bold;color:{couleur_decision};text-align:center;'>{texte_decision}</td>
        </tr>"""

    return f"""
    <table style='border-collapse:collapse;width:100%;table-layout:fixed;border:1px solid #c8d0dc;margin:10px 0 0 0;'>
        <thead>
            <tr style="background-color:#0D2E6E;height:30px;">
                <th style="padding:8px 10px;text-align:center;font-size:13px;font-weight:700;color:#ffffff;font-family:Arial,sans-serif;border-right:1px solid rgba(255,255,255,0.25);width:5%;">N°</th>
                <th style="padding:8px 10px;text-align:center;font-size:13px;font-weight:700;color:#ffffff;font-family:Arial,sans-serif;border-right:1px solid rgba(255,255,255,0.25);width:30%;">Élève</th>
                <th style="padding:8px 10px;text-align:center;font-size:13px;font-weight:700;color:#ffffff;font-family:Arial,sans-serif;border-right:1px solid rgba(255,255,255,0.25);width:10%;">Classe</th>
                <th style="padding:8px 10px;text-align:center;font-size:13px;font-weight:700;color:#ffffff;font-family:Arial,sans-serif;border-right:1px solid rgba(255,255,255,0.25);width:12%;">École</th>
                <th style="padding:8px 10px;text-align:center;font-size:13px;font-weight:700;color:#ffffff;font-family:Arial,sans-serif;border-right:1px solid rgba(255,255,255,0.25);width:30%;">Réf. Demande</th>
                <th style="padding:8px 10px;text-align:center;font-size:13px;font-weight:700;color:#ffffff;font-family:Arial,sans-serif;width:13%;">Décision</th>
            </tr>
        </thead>
        <tbody style='border-bottom:1px solid #c8d0dc;'>{lignes}</tbody>
    </table>
    <div style='height:20px;line-height:20px;font-size:1px;'>&nbsp;</div>
    """


# Fonction remplacement balises HTML
def remplacer_balises_html(doc, balises):
    for para in doc.paragraphs:
        for run in para.runs:
            for balise, valeur in balises.items():
                if balise in run.text:
                    run.text = run.text.replace(balise, valeur)

    lignes_html = []
    for i, para in enumerate(doc.paragraphs):
        texte = para.text

        # Gras sur NOM_FAMILLE et chere famille
        if texte.strip().startswith("Chère Famille"):
            texte = f"<strong>{texte}</strong>"

        nom = balises.get("{{NOM_FAMILLE}}", "")
        if nom and nom in texte:
            texte = texte.replace(nom, f"<b>{nom}</b>")

        # Gras sur email
        # Gras + couleur sur email
        email_val = balises.get("{{email}}", "")
        if email_val and email_val in texte:
            texte = texte.replace(
                email_val,
                f"<span style='font-weight:bold;color:#2F80ED;'>{email_val}</span>"
            )
        # Gras sur Matricule
        matricule_val = balises.get("{{Matricule}}", "")
        if matricule_val and matricule_val in texte and "matricule" in texte.lower():
            texte = texte.replace(matricule_val, f"<b>{matricule_val}</b>")

        # Tableau LISTE_ADMIS (tableau unique fusionné)
        liste_admis = balises.get("{{LISTE_ADMIS}}", "")
        if liste_admis and liste_admis in texte:
            texte = texte.replace(liste_admis, liste_admis)

        if "email de contact" in texte.lower():
            lignes_html.append(
                f"<p style='font-family:Arial,sans-serif;font-size:15px;color:#333;line-height:1.8;margin:0 0 0 0;'>{texte}</p>"
            )
            continue

        if "matricule famille" in texte.lower():
            lignes_html.append(
                f"<p style='font-family:Arial,sans-serif;font-size:15px;color:#333;line-height:1.8;margin:0 0 12px 0;'>{texte}</p>"
            )
            continue

        # Mettre la formule de politesse en gras
        if texte.strip() == "Cordialement,":
            texte = "<strong>Cordialement,</strong>"
            lignes_html.append(
                f"<p style='font-family:Arial,sans-serif;font-size:15px;color:#333;line-height:1.8;margin:0 0 0 0;'>{texte}</p>"
            )
            continue
        if texte.strip() in ["L'équipe de la coordination.", "L\u2019équipe de la coordination."]:
            texte = "<strong>L'équipe de la coordination.</strong>"
            lignes_html.append(
                f"<p style='font-family:Arial,sans-serif;font-size:15px;color:#333;line-height:1.8;margin:0 0 12px 0;'>{texte}</p>"
            )
            continue

        # Formalités
        if texte.strip() == "":
            continue

        if i == 0:
            lignes_html.append(
                f"<p style='font-family:Arial,sans-serif;font-size:16px;color:#1a1a2e;font-weight:bold;margin:0 0 16px 0;'>{texte}</p>"
            )
        elif "ci-après vos informations" in texte.lower():
            lignes_html.append(
                f"<p style='font-family:Arial,sans-serif;font-size:15px;color:#333;line-height:1.8;margin:20px 0 12px 0;'>{texte}</p>"
            )
        else:
            lignes_html.append(
                f"<p style='font-family:Arial,sans-serif;font-size:15px;color:#333;line-height:1.8;margin:0 0 12px 0;'>{texte}</p>"
            )
    # ajout d'image de signature

    signature_html = """
    <img src="cid:signature" width="600" height="200" style="display:block;align="center">
    """
    lignes_html.append(signature_html)
    return "<html><body style='font-family:Arial,sans-serif;font-size:15px;color:#333;margin:20px;'>" + "".join(lignes_html) + "</body></html>"

# Fonction pour générer un PDF avec ReportLab
def generer_pdf(groupe, output_path):

    # Marges réduites pour occuper toute la largeur
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=15,
        rightMargin=15,
        topMargin=25,
        bottomMargin=200
    )

    elements = []
    styles = getSampleStyleSheet()

    # Logo
    logo = Image(
        r"C:\Users\user\PycharmProjects\PythonProject\logo famille et education.png",
        width=150,
        height=45,
    )

    titre_style = ParagraphStyle(
        name="TitreEntete",
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=20,
        textColor=colors.HexColor("#1F4E79"),
        alignment=0,
        spaceBefore=0,
        spaceAfter=0
    )

    titre = Paragraph(
        "Résultats des délibérations - Demandes Inscriptions 2026-2027",
        titre_style
    )

    # Logo + titre sur la même ligne
    entete = Table(
        [[logo, titre]],
        colWidths=[105, 435]
    )

    entete.setStyle(TableStyle([

        # Alignement vertical
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),

        # Logo à gauche
        ('ALIGN', (0, 0), (0, 0), 'LEFT'),

        # Titre à gauche
        ('ALIGN', (1, 0), (1, 0), 'LEFT'),

        # Aucun espace autour du logo
        ('LEFTPADDING', (0, 0), (0, 0), 0),
        ('RIGHTPADDING', (0, 0), (0, 0), 0),

        # Espace entre le logo et le titre
        ('LEFTPADDING', (1, 0), (1, 0), 18),

        ('RIGHTPADDING', (1, 0), (1, 0), 0),

        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),

    ]))

    elements.append(entete)
    elements.append(Spacer(1, 10))

    # En-tête du tableau
    data = [[
        Paragraph("<font color='white'><b>N°</b></font>", styles["BodyText"]),
        Paragraph("<font color='white'><b>Élève</b></font>", styles["BodyText"]),
        Paragraph("<font color='white'><b>Classe</b></font>", styles["BodyText"]),
        Paragraph("<font color='white'><b>École</b></font>", styles["BodyText"]),
        Paragraph("<font color='white'><b>Réf. Demande</b></font>", styles["BodyText"]),
        Paragraph("<font color='white'><b>Décision</b></font>", styles["BodyText"])
    ]]

    # Lignes du tableau
    for i, row in enumerate(groupe.itertuples(), start=1):
        decision_val = str(row.decision).strip().lower()
        decision_texte = "Admis" if decision_val == "admis" else "Refusé"
        couleur = "green" if decision_val == "admis" else "red"

        data.append([
            Paragraph(str(i), styles["BodyText"]),
            Paragraph(f"<b>{row.nom_prenom_eleve}</b>", styles["BodyText"]),
            Paragraph(str(row.classe_sollicitee), styles["BodyText"]),
            Paragraph(str(row.ecole_sollicitee), styles["BodyText"]),
            Paragraph(str(row.ref_demande), styles["BodyText"]),
            Paragraph(f"<font color='{couleur}'><b>{decision_texte}</b></font>", styles["BodyText"])
        ])

    # Largeurs ajustées pour occuper toute la largeur A4

    elements.append(Spacer(1, 40))

    # Tableau centré
    table = Table(
        data,
        colWidths=[20, 150, 45, 60, 160, 55],
        repeatRows=1,
        hAlign="CENTER"
    )

    style = [
        # En-tête bleu de logo + texte blanc
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#1873AB")),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),

        # Corps du tableau avec fond uniforme
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor("#F0F8FF")),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('ALIGN', (0, 1), (0, -1), 'CENTER'),
        ('ALIGN', (2, 1), (3, -1), 'CENTER'),
        ('ALIGN', (5, 1), (5, -1), 'CENTER'),
        ('VALIGN', (0, 1), (-1, -1), 'MIDDLE'),

        # Bordures fines pour former les cellules
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),

        # Padding
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),

        # Wordwrap pour éviter débordement
        ('WORDWRAP', (0, 0), (-1, -1), 'CJK')
    ]

    # Couleur de la décision
    for ligne in range(1, len(data)):
        decision = str(groupe.iloc[ligne - 1]["decision"]).strip().lower()
        if decision == "admis":
            style.append(('TEXTCOLOR', (5, ligne), (5, ligne), colors.green))
            style.append(('FONTNAME', (5, ligne), (5, ligne), 'Helvetica-Bold'))
        else:
            style.append(('TEXTCOLOR', (5, ligne), (5, ligne), colors.red))
            style.append(('FONTNAME', (5, ligne), (5, ligne), 'Helvetica-Bold'))

    table.setStyle(TableStyle(style))
    elements.append(table)

    # Espace avant le footer
    elements.append(Spacer(1, 40))

    # Trait horizontal
    elements.append(HRFlowable(width="100%", thickness=0.5, color=colors.blue))

    # Date du jour
    date_du_jour = datetime.now().strftime("%d/%m/%Y")
    date_style = ParagraphStyle(
        name="DateBasPage",
        fontName="Helvetica",
        fontSize=10,
        alignment=2,
        textColor=colors.orange
    )
    elements.append(Paragraph(f"Date d'envoi : {date_du_jour}", date_style))
    doc.build(elements)
# Textes fixes
INTRO_ADMIS = "Nous avons le plaisir de vous informer que votre dossier de demandes d'inscription soumis pour l'année scolaire 2026-2027 a été étudié avec attention et a reçu un avis favorable pour l'ensemble des demandes.\nCi-après les détails :"
INTRO_MIXTE = "Nous avons le plaisir de vous informer que votre dossier de demandes d'inscription soumis pour l'année scolaire 2026-2027 a été étudié avec attention et a reçu un avis favorable pour au moins une demande. \nCi-après les détails : "
INTRO_REFUS_TOTAL = """Après étude attentive de votre dossier, nous regrettons de vous informer que nous ne sommes pas en mesure d'y donner une suite favorable à ce jour.
                        Cette décision peut être liée à des contraintes d'effectifs et/ou à d'autres critères d'admission définis par l'établissement."""

FORMALITES = """Tout en vous encourageant à prendre en compte les délais indiqués par l'établissement, nous vous invitons à procéder aux formalités administratives et financières nécessaires à la finalisation du processus d'inscription, notamment :
<ul style='font-family:Arial,sans-serif;font-size:15px;color:#333;line-height:1.4;margin:8px 0 8px 0;padding-left:20px;'>
  <li style='margin-bottom:4px;'>Engagement des parents pour suivre l'esprit du projet</li>
  <li style='margin-bottom:4px;'>Paiement de la réservation (si pas encore fait)</li>
  <li style='margin-bottom:4px;'>Chargement des dossiers</li>
</ul>
"""
Formalite_refus = """ Tout en vous souhaitant une issue favorable pour la prochaine échéance, nous vous remercions pour l'intérêt porté à notre établissement."""

# Traitement par famille
familles = df.groupby(["nom_famille", "email"])
MODE_TEST = True
outlook = win32.Dispatch("Outlook.Application")

for (nom_famille, email_parent), groupe in familles:

    admis = groupe[groupe["decision"] == "Admis"]
    refuses = groupe[groupe["decision"] == "Refuse"]

    # Un seul tableau fusionné contenant TOUS les enfants (admis + refusés)

    tableau_fusionne = construire_tableau_html(groupe)

    if len(admis) > 0 and len(refuses) == 0:
        print(f"Cas TOUS ADMIS : {nom_famille}")
        doc = Document("modele_admis.docx")
        balises = {
            "{{NOM_FAMILLE}}": nom_famille,
            "{{INTRODUCTION}}": INTRO_ADMIS,
            "{{LISTE_ADMIS}}": tableau_fusionne,
            "{{FORMALITE}}": FORMALITES,
            "{{Matricule}}": groupe["Matricule"].iloc[0],
            "{{email}}": email_parent,
        }

    elif len(admis) > 0 and len(refuses) > 0:
        print(f"Cas MIXTE : {nom_famille}")
        doc = Document("modele_admis.docx")
        balises = {
            "{{NOM_FAMILLE}}": nom_famille,
            "{{INTRODUCTION}}": INTRO_MIXTE,
            "{{LISTE_ADMIS}}": tableau_fusionne,
            "{{FORMALITE}}": FORMALITES,
            "{{Matricule}}": groupe["Matricule"].iloc[0],
            "{{email}}": email_parent,
        }

    else:
        print(f"Cas AUCUN ADMIS : {nom_famille}")
        doc = Document("modele_refus.docx")
        balises = {
            "{{NOM_FAMILLE}}": nom_famille,
            "{{INTRODUCTION}}": INTRO_REFUS_TOTAL,
            "{{LISTE_REFUS}}": tableau_fusionne,
            "{{FORMALITE}}": Formalite_refus,
        }

    corps_html = remplacer_balises_html(doc, balises)

    # Générer le PDF pour cette famille

    pdf_path = f"C:\\Users\\user\\PycharmProjects\\PythonProject\\Résultats des délibérations - Demandes Inscriptions 2026-2027.pdf"
    generer_pdf(groupe, pdf_path)

    # Création du mail Outlook

    def get_account_by_email(outlook, email):
        for account in outlook.Session.Accounts:
            if account.SmtpAddress.lower() == email.lower():
                return account
        raise Exception(f"Compte {email} introuvable dans Outlook.")


    # Exemple d’utilisation

    COMPTE_EXPEDITEUR = "infos-coordination@lecoledesfamilles.org"
    compte_cible = get_account_by_email(outlook, COMPTE_EXPEDITEUR)

    mail = outlook.CreateItem(0)

    mail.SendUsingAccount = compte_cible

    # Destinataire

    mail.To = email_parent

    mail.Subject = "Résultats des délibérations - Demandes Inscriptions 2026-2027"

    mail.HTMLBody = corps_html

    # Signature

    attachment = mail.Attachments.Add(
        r"C:\Users\user\PycharmProjects\PythonProject\signature.png"
    )

    attachment.PropertyAccessor.SetProperty(
        "http://schemas.microsoft.com/mapi/proptag/0x3712001F",
        "signature"
    )
    # PDF

    mail.Attachments.Add(pdf_path)
    # Envoi
    MODE_TEST = True
    if MODE_TEST:

        mail.Display()

        print(
            f"Mail préparé pour : {email_parent} ({nom_famille})"
        )

    else:

        mail.Send()

        print(
            f"Mail envoyé à : {email_parent} ({nom_famille})"
        )

print("Terminé !")

